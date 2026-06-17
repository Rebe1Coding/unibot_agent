"""Shared DOCX builder -- Markdown to GOST-formatted Word document.

Used by both the agent tool (university-agent) and the Celery worker.
"""

from __future__ import annotations

import io
import re

# -- GOST formatting constants ------------------------------------------------

# GOST 7.32-2017 & GOST 2.105-2019
FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14
FONT_SIZE_HEADING1_PT = 16
FONT_SIZE_HEADING2_PT = 14
LINE_SPACING = 1.5
MARGIN_LEFT_CM = 3.0
MARGIN_RIGHT_CM = 1.0
MARGIN_TOP_CM = 2.0
MARGIN_BOTTOM_CM = 2.0
FIRST_LINE_INDENT_CM = 1.25


def markdown_to_docx(md: str, depersonalize: bool = True) -> bytes:
    """Convert Markdown text to GOST-formatted .docx bytes.

    Args:
        md: Markdown source text.
        depersonalize: If True, replace first-person constructions with
                       impersonal ones (GOST requirement for scientific text).

    Returns:
        Bytes of the generated .docx file.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # -- Page setup (GOST margins) ---------------------------------------------
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)

    # -- Default paragraph style -----------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_PT)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING
    pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)

    # -- Configure heading styles ----------------------------------------------
    for level, size in [(1, FONT_SIZE_HEADING1_PT), (2, FONT_SIZE_HEADING2_PT)]:
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = FONT_NAME
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        hstyle.paragraph_format.space_before = Pt(12)
        hstyle.paragraph_format.space_after = Pt(6)
        hstyle.paragraph_format.first_line_indent = Cm(0)
        hstyle.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        )

    # -- Parse and render Markdown ---------------------------------------------
    lines = md.split("\n")
    i = 0
    list_counter = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            style_name = f"Heading {min(level, 2)}"
            p = doc.add_paragraph(style=style_name)
            _add_formatted_run(p, text)
            list_counter = 0
            i += 1
            continue

        # Unordered list
        if re.match(r"^[-*+]\s+", stripped):
            text = re.sub(r"^[-*+]\s+", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(FIRST_LINE_INDENT_CM)
            _add_formatted_run(p, f"\u2013 {text}")
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if ol_match:
            list_counter += 1
            text = ol_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(FIRST_LINE_INDENT_CM)
            _add_formatted_run(p, f"{list_counter}) {text}")
            i += 1
            continue

        # Regular paragraph
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"^(#{1,3}\s|[-*+]\s|\d+[.)]\s)", nxt):
                break
            para_lines.append(nxt)
            i += 1

        text = " ".join(para_lines)
        if depersonalize:
            text = _depersonalize(text)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_formatted_run(p, text)
        list_counter = 0

    # -- Page numbering (centered bottom, GOST requirement) --------------------
    _add_page_numbers(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_run(paragraph, text: str) -> None:
    """Parse inline Markdown (bold, italic) and add runs."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)


def _depersonalize(text: str) -> str:
    """Replace first-person constructions with impersonal ones (GOST requirement)."""
    replacements = [
        (r"\bЯ считаю\b", "Считается"),
        (r"\bя считаю\b", "считается"),
        (r"\bМы считаем\b", "Считается"),
        (r"\bмы считаем\b", "считается"),
        (r"\bЯ полагаю\b", "Полагается"),
        (r"\bя полагаю\b", "полагается"),
        (r"\bМы полагаем\b", "Полагается"),
        (r"\bмы полагаем\b", "полагается"),
        (r"\bЯ предлагаю\b", "Предлагается"),
        (r"\bя предлагаю\b", "предлагается"),
        (r"\bМы предлагаем\b", "Предлагается"),
        (r"\bмы предлагаем\b", "предлагается"),
        (r"\bЯ использую\b", "Используется"),
        (r"\bя использую\b", "используется"),
        (r"\bМы используем\b", "Используется"),
        (r"\bмы используем\b", "используется"),
        (r"\bЯ рассматриваю\b", "Рассматривается"),
        (r"\bя рассматриваю\b", "рассматривается"),
        (r"\bМы рассматриваем\b", "Рассматривается"),
        (r"\bмы рассматриваем\b", "рассматривается"),
        (r"\bЯ провёл\b", "Было проведено"),
        (r"\bя провёл\b", "было проведено"),
        (r"\bМы провели\b", "Было проведено"),
        (r"\bмы провели\b", "было проведено"),
        (r"\bЯ выполнил\b", "Было выполнено"),
        (r"\bя выполнил\b", "было выполнено"),
        (r"\bМы выполнили\b", "Было выполнено"),
        (r"\bмы выполнили\b", "было выполнено"),
        (r"\bЯ разработал\b", "Было разработано"),
        (r"\bя разработал\b", "было разработано"),
        (r"\bМы разработали\b", "Было разработано"),
        (r"\bмы разработали\b", "было разработано"),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)
    return text


def _add_page_numbers(doc) -> None:
    """Add centered page numbers to the footer (GOST 7.32-2017)."""
    from docx.oxml.ns import qn

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1  # CENTER

        run = p.add_run()
        fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char_begin)

        run2 = p.add_run()
        instr = run2._element.makeelement(qn("w:instrText"), {})
        instr.text = " PAGE "
        run2._element.append(instr)

        run3 = p.add_run()
        fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fld_char_end)
